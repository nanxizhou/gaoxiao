# 案例 1 ：批量生成法务函
# 目标：把【封号名单.xlsx】工作簿中的每个封号人员的名字和微信号添加到【法务函模板.docx】Word 文件对应的位置上，
# 并将【法务函模板.docx】Word 文件另存为【法务函-XXX.docx】 Word 文件
# 请随时查看知识库和案例练习助手，与自己编写代码的步骤和内容比对参考，训练思维
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt

wb = load_workbook('../工作/封号名单.xlsx')

ws = wb.active
i = 0
for row in ws.iter_rows(min_row=2,  max_row=26,  min_col=1,  max_col=2,  values_only=True):
    # 姓名
    name = row[0]
    # 微信号
    wxid = row[1]
    # 获取Document
    doc = Document('../工作/法务函模板.docx')
    # 获取段落
    para = doc.paragraphs[5]
    # 添加run对象
    # run = para.add_run()
    # 添加封号人员的姓名
    run_name = para.add_run(name)
    # 加粗
    run_name.font.bold = True
    # 下划线
    run_name.font.underline = True
    # 字体大小
    run_name.font.size = Pt(14)
    # 添加封号人员的微信
    run_wxid = para.add_run('同学（WeChat ID: {}）'.format(wxid))
    # 加粗
    run_wxid.font.bold = True
    # 下划线
    run_wxid.font.underline = True
    # 字体大小
    run_wxid.font.size = Pt(14)
    i += 1
    doc.save('../法务/{}法务函.docx'.format(name))
